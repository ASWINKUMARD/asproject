from flask import Flask, render_template, request, jsonify, send_file, make_response
from flask_cors import CORS
from dotenv import load_dotenv
load_dotenv()

import os
import logging
import requests
import shutil
import pickle
from typing import Any, List, Optional, Dict
import re
import io
import uuid
from datetime import datetime
import json
from werkzeug.utils import secure_filename

# For document generation
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from langchain_community.document_loaders import (
    PyPDFLoader, Docx2txtLoader, UnstructuredExcelLoader,
    UnstructuredCSVLoader, TextLoader, UnstructuredHTMLLoader,
    UnstructuredMarkdownLoader
)
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.language_models.llms import LLM
from langchain_core.embeddings import Embeddings
from langchain_core.documents import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
OPENROUTER_API_BASE = "https://openrouter.ai/api/v1/chat/completions"
MODEL_NAME = "kwaipilot/kat-coder-pro:free"

DATA_DIR = "./data"
FAISS_DB_PATH = "./faiss_db"
CHAT_HISTORY_DIR = "./chat_history"

os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(FAISS_DB_PATH, exist_ok=True)
os.makedirs(CHAT_HISTORY_DIR, exist_ok=True)

SUPPORTED_EXTENSIONS = {
    ".pdf": PyPDFLoader,
    ".docx": Docx2txtLoader,
    ".xlsx": UnstructuredExcelLoader,
    ".xls": UnstructuredExcelLoader,
    ".csv": UnstructuredCSVLoader,
    ".txt": TextLoader,
    ".html": UnstructuredHTMLLoader,
    ".htm": UnstructuredHTMLLoader,
    ".md": UnstructuredMarkdownLoader,
}

llm = None
qa_chain = None
vector_db = None
embeddings = None
chat_history: Dict[str, Dict[str, Any]] = {}

app = Flask(__name__)
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 150 * 1024 * 1024

def create_chat(user_msg: str, bot_response: str, is_table: bool = False) -> Dict[str, Any]:
    chat_id = str(uuid.uuid4())
    now = datetime.utcnow().isoformat() + "Z"
    
    # Create a chat with conversation history to support memory
    chat_entry = {
        "id": chat_id,
        "messages": [{"role": "user", "content": user_msg}, {"role": "assistant", "content": bot_response}],
        "is_table": is_table,
        "created_at": now,
        "updated_at": now,
    }
    
    chat_history[chat_id] = chat_entry
    _save_chat_to_file(chat_id, chat_entry)
    return chat_entry

def update_chat(chat_id: str, bot_response: Optional[str] = None, user_message: Optional[str] = None) -> Optional[Dict[str, Any]]:
    chat = chat_history.get(chat_id)
    if not chat:
        return None
    
    # Initialize messages list if it doesn't exist
    if "messages" not in chat:
        chat["messages"] = []
    
    # Add user message if provided
    if user_message is not None:
        chat["messages"].append({"role": "user", "content": user_message})
    
    # Add bot response if provided
    if bot_response is not None:
        chat["messages"].append({"role": "assistant", "content": bot_response})
    
    chat["updated_at"] = datetime.utcnow().isoformat() + "Z"
    _save_chat_to_file(chat_id, chat)
    return chat

def delete_chat(chat_id: str) -> bool:
    if chat_id in chat_history:
        del chat_history[chat_id]
        filepath = os.path.join(CHAT_HISTORY_DIR, f"{chat_id}.json")
        if os.path.exists(filepath):
            os.remove(filepath)
        return True
    return False

def _save_chat_to_file(chat_id: str, chat_entry: Dict[str, Any]):
    filepath = os.path.join(CHAT_HISTORY_DIR, f"{chat_id}.json")
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(chat_entry, f, ensure_ascii=False, indent=2)

def _load_chat_history():
    if not os.path.exists(CHAT_HISTORY_DIR):
        return
    for filename in os.listdir(CHAT_HISTORY_DIR):
        if filename.endswith(".json"):
            filepath = os.path.join(CHAT_HISTORY_DIR, filename)
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    chat_entry = json.load(f)
                    
                    # Migrate old format to new format if needed
                    if "messages" not in chat_entry:
                        # Convert old format to new format
                        messages = []
                        
                        # Add user message if it exists
                        if "user_message" in chat_entry:
                            messages.append({"role": "user", "content": chat_entry["user_message"]})
                        
                        # Add bot responses if they exist
                        if "bot_responses" in chat_entry:
                            for response in chat_entry["bot_responses"]:
                                messages.append({"role": "assistant", "content": response})
                        elif "bot_response" in chat_entry:
                            # Handle single bot response
                            messages.append({"role": "assistant", "content": chat_entry["bot_response"]})
                        
                        # Create new format
                        new_chat_entry = {
                            "id": chat_entry["id"],
                            "messages": messages,
                            "is_table": chat_entry.get("is_table", False),
                            "created_at": chat_entry.get("created_at", datetime.utcnow().isoformat() + "Z"),
                            "updated_at": chat_entry.get("updated_at", datetime.utcnow().isoformat() + "Z")
                        }
                        
                        chat_history[chat_entry["id"]] = new_chat_entry
                        
                        # Save back in new format
                        _save_chat_to_file(chat_entry["id"], new_chat_entry)
                    else:
                        # Chat is already in new format
                        chat_history[chat_entry["id"]] = chat_entry
            except Exception as e:
                logger.error(f"Failed to load chat history file {filename}: {e}")

def detect_table_request(user_msg: str) -> bool:
    table_keywords = [
        "table", "tabular", "spreadsheet", "rows and columns",
        "in a table", "as a table", "table format", "create table"
    ]
    
    # Only return True if user explicitly requests a table format
    # Never generate tables for any other requests
    has_table_keyword = any(keyword in user_msg.lower() for keyword in table_keywords)
    
    return has_table_keyword

# Table-related functions have been removed as per requirement to eliminate all table formats

# The following functions were removed:
# - parse_table_from_response
# - convert_to_4column_format
# - extract_category_from_item
# - format_clean_table
# - add_totals_to_table


def identify_service_type(service_text: str) -> str:
    """Properly identify the service type from text, prioritizing design services over mobile services."""
    service_lower = service_text.lower().strip()
    
    # Check for design-related keywords first (higher priority)
    if any(keyword in service_lower for keyword in ['ui/ux', 'ui', 'ux', 'design']):
        if 'mobile' in service_lower or 'app' in service_lower:
            # This is a mobile app UI/UX design service
            if not service_lower.endswith(('design', 'development', 'service', 'app', 'bot', 'system', 'solution', 'project')):
                return service_text.strip().title() + " Design"
            return service_text.strip().title()
        elif not service_lower.endswith(('design', 'development', 'service', 'app', 'bot', 'system', 'solution', 'project')):
            return service_text.strip().title() + " Design"
        return service_text.strip().title()
    
    # Check for other specific service types
    if service_lower == "ios":
        return "iOS Development"
    elif service_lower == "android":
        return "Android Development"
    elif service_lower == "web":
        return "Website Development"
    elif service_lower == "mobile":
        return "Mobile App Development"
    
    # Default handling
    if not service_lower.endswith(('development', 'service', 'app', 'bot', 'system', 'solution', 'project', 'design')):
        # Check if it's a design service
        if any(keyword in service_lower for keyword in ['design', 'ui', 'ux']):
            return service_text.strip().title() + " Design"
        # Check if it's a development service
        elif any(keyword in service_lower for keyword in ['development', 'dev']):
            return service_text.strip().title() + " Development"
        # Default to Development
        else:
            return service_text.strip().title() + " Development"
    
    return service_text.strip().title()

def get_tech_stack_for_service(service_name: str) -> str:
    """Return a generic tech stack as the actual stack should be retrieved from the documents."""
    # Since the actual tech stack should come from the retrieved documents,
    # we return a placeholder that indicates the information should be retrieved from context
    return "Technology stack information will be retrieved from the provided documents based on the specific service requirements.\nPlease refer to the retrieved context for detailed technology stack information."

def remove_markdown_formatting(text: str) -> str:
    """Remove markdown formatting symbols from text."""
    if not text:
        return text
    
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'__(.*?)__', r'\1', text)
    
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'``.*?```', '', text, flags=re.DOTALL)
    
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    text = re.sub(r'~~(.*?)~~', r'\1', text)
    
    return text.strip()

def generate_comprehensive_format(service_name, original_response: str, user_request: str = "") -> str:
    """Generate comprehensive format with service name, pricing, explanation, and tech stack in conversational style."""
    # Extract service-specific information from the original response
    service_info = original_response[:500] + "..." if len(original_response) > 500 else original_response
    
    import re
    
    # If user_request is provided, extract price and quantity from the user's request
    unit_price = None
    quantity = 1
    
    # Check if the user wants to add multiple services to a quotation
    multiple_services = []
    
    if user_request:
        user_lower = user_request.lower()
        
        # Check if user is asking for multiple services in a single quotation
        # First, look for services mentioned with 'for' keyword - extract everything after 'for'
        # This handles "create quotation for Website AI Chatbot and Android App Development"
        for_pattern = re.search(r'for\s+(.+?)(?:$|quotation|with|and|\+)', user_request, re.IGNORECASE)
        if for_pattern:
            services_text = for_pattern.group(1).strip()
            
            # Split by 'and', '&', 'with', or '+' to get individual services
            service_parts = []
            
            # First try to split by 'and' (but avoid splitting 'Android')
            # Split by ' and ' (with spaces) but not 'android'
            parts = re.split(r'\s+and\s+|\s*&\s*|\s+with\s+|\s+\+\s+', services_text)
            
            for part in parts:
                part = part.strip()
                if part:
                    # Enhance common short forms
                    service_lower = part.lower().strip()
                    if service_lower == "ios":
                        part = "iOS Development"
                    elif service_lower == "android":
                        part = "Android Development"  # Keep Android as a complete service name
                    elif service_lower == "web":
                        part = "Website Development"
                    elif service_lower == "mobile":
                        part = "Mobile App Development"
                    elif not any(s in service_lower for s in ["android"]):  # Don't modify if it contains 'android'
                        # Capitalize properly
                        part = part.strip().title()
                        
                        # Add 'Development' suffix if not already present and seems appropriate
                        if not part.lower().endswith(('development', 'service', 'app', 'bot', 'system', 'solution', 'project')) and len(part.split()) <= 4:
                            if 'chatbot' in part.lower() or 'ai' in part.lower():
                                # For AI Chatbot, keep as is
                                pass
                            elif 'website' in part.lower() or 'web' in part.lower():
                                part = 'Website Development'
                            elif 'android' in part.lower():
                                part = 'Android Development'
                            elif 'ios' in part.lower():
                                part = 'iOS Development'
                            else:
                                part = part + " Development"
                    
                    if part not in multiple_services:
                        multiple_services.append(part)
        
        # Also look for other patterns like "add X to quotation"
        add_include_patterns = [
            r'add\s+([a-zA-Z\s]+?)\s+(?:to|in|with)\s+quotation',  # "add iOS Development to quotation"
            r'include\s+([a-zA-Z\s]+?)\s+(?:to|in|with)\s+quotation',  # "include Android Development in quotation"
        ]
        
        for pattern in add_include_patterns:
            matches = re.findall(pattern, user_lower, re.IGNORECASE)
            for match in matches:
                service = match.strip()
                if service:
                    # Enhance common short forms
                    service_lower = service.lower().strip()
                    if service_lower == "ios":
                        service = "iOS Development"
                    elif service_lower == "android":
                        service = "Android Development"
                    elif service_lower == "web":
                        service = "Website Development"
                    elif service_lower == "mobile":
                        service = "Mobile App Development"
                    else:
                        service = service.title()
                        if not service.lower().endswith(('development', 'service', 'app', 'bot', 'system', 'solution', 'project')):
                            service = service + " Development"
                    if service not in multiple_services:
                        multiple_services.append(service)
        

        
        # Extract user-specified price
        price_patterns = [
            r'rs\.?\s*([0-9,]+)',  # rs 50000 or rs. 50000
            r'₹\s*([0-9,]+)',      # ₹50,000
            r'price\s+([0-9,]+)',   # price 50000
        ]
        
        for pattern in price_patterns:
            price_match = re.search(pattern, user_lower)
            if price_match:
                try:
                    unit_price = int(price_match.group(1).replace(',', ''))
                    break
                except ValueError:
                    continue
        
        # Extract user-specified quantity
        quantity_patterns = [
            r'quantity\s*-?\s*(\d+)',  # quantity - 2 or quantity 2
            r'qty\s*-?\s*(\d+)',      # qty - 2 or qty 2
            r'x\s*(\d+)',              # x 2
            r'(\d+)\s*x',              # 2 x
        ]
        
        for pattern in quantity_patterns:
            qty_match = re.search(pattern, user_lower)
            if qty_match:
                try:
                    quantity = int(qty_match.group(1))
                    break
                except ValueError:
                    continue
        
        # If both unit price and quantity are specified by user, calculate total
        if unit_price is not None:
            total_price = unit_price * quantity
            price = f'₹{total_price:,} (₹{unit_price:,} × {quantity})'
        else:
            # Fall back to extracting from original response
            price_patterns = [
                r'₹\s*([0-9,]+)',  # ₹1,00,000
                r'Rs\.?\s*([0-9,]+)',  # Rs. 1,00,000 or Rs 1,00,000
                r'([0-9,]{5,})'  # Standalone large numbers (likely prices)
            ]
            
            price_matches = []
            for pattern in price_patterns:
                matches = re.findall(pattern, original_response)
                price_matches.extend(matches)
            
            # Remove duplicates while preserving order
            seen = set()
            unique_matches = []
            for match in price_matches:
                if match not in seen:
                    seen.add(match)
                    unique_matches.append(match)
            
            if unique_matches:
                # Use the first extracted price (most likely the main price)
                clean_price = int(unique_matches[0].replace(',', ''))
                price = f'₹{clean_price:,}'
            else:
                price = "₹1,00,000"  # Default price
    else:
        # Extract from original response only
        price_patterns = [
            r'₹\s*([0-9,]+)',  # ₹1,00,000
            r'Rs\.?\s*([0-9,]+)',  # Rs. 1,00,000 or Rs 1,00,000
            r'([0-9,]{5,})'  # Standalone large numbers (likely prices)
        ]
        
        price_matches = []
        for pattern in price_patterns:
            matches = re.findall(pattern, original_response)
            price_matches.extend(matches)
        
        # Remove duplicates while preserving order
        seen = set()
        unique_matches = []
        for match in price_matches:
            if match not in seen:
                seen.add(match)
                unique_matches.append(match)
        
        if unique_matches:
            # Use the first extracted price (most likely the main price)
            clean_price = int(unique_matches[0].replace(',', ''))
            price = f'₹{clean_price:,}'
        else:
            price = "₹1,00,000"  # Default price
    
    # Check if service_name is a list (multiple services) or a single string
    if isinstance(service_name, list):
        # Handle multiple services directly passed to the function
        service_descriptions = []
        total_cost = 0
        
        # Generate individual cost descriptions and calculate total
        individual_costs = []
        total_cost = 0  # Reset total cost for this section
        for serv_name in service_name:
            # First, try to extract specific price for this service from the original response
            # This allows RAG results to override default pricing
            clean_price = None
            
            # Look for service-specific pricing in the original response
            service_price_pattern = re.search(rf'{re.escape(serv_name)}.*?₹\s*([0-9,]+)|{re.escape(serv_name)}.*?Rs\.?\s*([0-9,]+)', original_response, re.IGNORECASE)
            if service_price_pattern:
                price_match = service_price_pattern.group(1) or service_price_pattern.group(2)
                if price_match:
                    try:
                        clean_price = int(price_match.replace(',', ''))
                    except ValueError:
                        clean_price = None
            
            # If no service-specific price found, look for any price mentioned in the response
            if clean_price is None:
                # Extract prices from the original response
                price_patterns = [
                    r'₹\s*([0-9,]+)',  # ₹1,00,000
                    r'Rs\.?\s*([0-9,]+)',  # Rs. 1,00,000 or Rs 1,00,000
                    r'([0-9,]{5,})'  # Standalone large numbers (likely prices)
                ]
                
                price_matches = []
                for pattern in price_patterns:
                    matches = re.findall(pattern, original_response)
                    price_matches.extend(matches)
                
                # Remove duplicates while preserving order
                seen = set()
                unique_matches = []
                for match in price_matches:
                    if match not in seen:
                        seen.add(match)
                        unique_matches.append(match)
                
                if unique_matches:
                    # Use the first extracted price (most likely the main price)
                    try:
                        clean_price = int(unique_matches[0].replace(',', ''))
                    except ValueError:
                        clean_price = None
            
            # If still no price found from the response, use service-specific default
            if clean_price is None:
                # Determine price based on the service type
                if 'ai chatbot' in serv_name.lower() or 'chatbot' in serv_name.lower():
                    clean_price = 45000
                elif 'payment' in serv_name.lower() or 'gateway' in serv_name.lower():
                    clean_price = 15000
                elif 'design' in serv_name.lower() or 'ui' in serv_name.lower() or 'ux' in serv_name.lower():
                    clean_price = 25000
                elif 'mobile' in serv_name.lower() or 'app' in serv_name.lower():
                    clean_price = 30000
                elif 'website' in serv_name.lower():
                    clean_price = 20000
                else:
                    # For services that don't match specific types, use a reasonable default
                    clean_price = 50000  # Default price for unmatched services
            
            individual_costs.append(f"For {serv_name}, Syngrid offers a fixed cost of ₹{clean_price:,}.")
            total_cost += clean_price  # Add to total cost
        
        service_overviews = []
        for serv_name in service_name:
            # Generate a generic service description that works for any service
            service_description = f"{serv_name} is a specialized service that involves creating custom solutions to meet specific business requirements. This service includes planning, development, testing, and deployment of tailored solutions to address unique challenges."
            
            service_overviews.append(f"Regarding {serv_name} specifically:\n\n{service_description}")
        
        comprehensive_content = f"Company:\n\nSyngrid is a software development company based in Madurai, India, operating from the TBI@TCE at Thiruparankundaram. The company specializes in a range of digital services and can be contacted at (+91) 89250-31568.\n\n\nService Overview:\n\n" + "\n\n".join(service_overviews) + f"\n\n\nCost:\n\n" + "\n".join(individual_costs) + f"\n\nThe combined total for all services would be ₹{total_cost:,}.\n\n\nTechnology Stack:\n\n"
        
        # Extract technology stack information from the original response (RAG results)
        # Look for technology stack related content in the original response
        import re
        tech_stack_lines = []
        
        # Look for common tech stack patterns in the original response
        lines = original_response.split('\n')
        for line in lines:
            line = line.strip()
            # Look for lines that might contain tech stack information
            if any(keyword in line.lower() for keyword in ['frontend', 'backend', 'database', 'cloud', 'framework', 'library', 'technology', 'tool', 'platform', 'api', 'development', 'programming', 'programming language', 'javascript', 'python', 'java', 'react', 'angular', 'node', 'mysql', 'mongodb', 'aws', 'azure', 'docker', 'kubernetes']):
                if len(line) > 10 and not line.startswith('Regarding') and not line.startswith('Company') and not line.startswith('Cost'):
                    tech_stack_lines.append(line)
        
        if tech_stack_lines:
            comprehensive_content += "\n".join(tech_stack_lines)
        else:
            # Fallback to a general message about retrieving from documents
            comprehensive_content += "Technology stack information will be retrieved from the provided documents based on the specific service requirements.\nPlease refer to the retrieved context for detailed technology stack information."
        
        return comprehensive_content
    elif multiple_services and len(multiple_services) > 1:
        # Generate content for multiple services
        service_descriptions = []
        total_cost = 0
        
        individual_costs = []
        total_cost = 0  
        for serv_name in multiple_services:
            # First, try to extract specific price for this service from the original response
            # This allows RAG results to override default pricing
            clean_price = None
            
            # Look for service-specific pricing in the original response
            service_price_pattern = re.search(rf'{re.escape(serv_name)}.*?₹\s*([0-9,]+)|{re.escape(serv_name)}.*?Rs\.?\s*([0-9,]+)', original_response, re.IGNORECASE)
            if service_price_pattern:
                price_match = service_price_pattern.group(1) or service_price_pattern.group(2)
                if price_match:
                    try:
                        clean_price = int(price_match.replace(',', ''))
                    except ValueError:
                        clean_price = None
            
            # If no service-specific price found, look for any price mentioned in the response
            if clean_price is None:
                # Extract prices from the original response
                price_patterns = [
                    r'₹\s*([0-9,]+)',  # ₹1,00,000
                    r'Rs\.?\s*([0-9,]+)',  # Rs. 1,00,000 or Rs 1,00,000
                    r'([0-9,]{5,})'  # Standalone large numbers (likely prices)
                ]
                
                price_matches = []
                for pattern in price_patterns:
                    matches = re.findall(pattern, original_response)
                    price_matches.extend(matches)
                
                # Remove duplicates while preserving order
                seen = set()
                unique_matches = []
                for match in price_matches:
                    if match not in seen:
                        seen.add(match)
                        unique_matches.append(match)
                
                if unique_matches:
                    # Use the first extracted price (most likely the main price)
                    try:
                        clean_price = int(unique_matches[0].replace(',', ''))
                    except ValueError:
                        clean_price = None
            
            # If still no price found from the response, use service-specific default
            if clean_price is None:
                # Determine price based on the service type
                if 'ai chatbot' in serv_name.lower() or 'chatbot' in serv_name.lower():
                    clean_price = 45000
                elif 'payment' in serv_name.lower() or 'gateway' in serv_name.lower():
                    clean_price = 15000
                elif 'design' in serv_name.lower() or 'ui' in serv_name.lower() or 'ux' in serv_name.lower():
                    clean_price = 25000
                elif 'mobile' in serv_name.lower() or 'app' in serv_name.lower():
                    clean_price = 30000
                elif 'website' in serv_name.lower():
                    clean_price = 20000
                else:
                    # For services that don't match specific types, use a reasonable default
                    clean_price = 50000  # Default price for unmatched services
            
            individual_costs.append(f"For {serv_name}, Syngrid offers a fixed cost of ₹{clean_price:,}.")
            total_cost += clean_price  # Add to total cost
        
        # Generate the comprehensive format for multiple services
        service_overviews = []
        for serv_name in multiple_services:
            # Generate a generic service description that works for any service
            service_description = f"{serv_name} is a specialized service that involves creating custom solutions to meet specific business requirements. This service includes planning, development, testing, and deployment of tailored solutions to address unique challenges."
            
            service_overviews.append(f"Regarding {serv_name} specifically:\n\n{service_description}")
        
        comprehensive_content = f"Company:\n\nSyngrid is a software development company based in Madurai, India, operating from the TBI@TCE (Technology Business Incubator) at Thiruparankundaram. The company specializes in a range of digital services and can be contacted at (+91) 89250-31568.\n\n\nService Overview:\n\n" + "\n\n".join(service_overviews) + f"\n\n\nCost:\n\n" + "\n".join(individual_costs) + f"\n\nThe combined total for all services would be ₹{total_cost:,}.\n\n\nTechnology Stack:\n\n"
        
        # Extract technology stack information from the original response (RAG results)
        # Look for technology stack related content in the original response
        import re
        tech_stack_lines = []
        
        # Look for common tech stack patterns in the original response
        lines = original_response.split('\n')
        for line in lines:
            line = line.strip()
            # Look for lines that might contain tech stack information
            if any(keyword in line.lower() for keyword in ['frontend', 'backend', 'database', 'cloud', 'framework', 'library', 'technology', 'tool', 'platform', 'api', 'development', 'programming', 'programming language', 'javascript', 'python', 'java', 'react', 'angular', 'node', 'mysql', 'mongodb', 'aws', 'azure', 'docker', 'kubernetes']):
                if len(line) > 10 and not line.startswith('Regarding') and not line.startswith('Company') and not line.startswith('Cost'):
                    tech_stack_lines.append(line)
        
        if tech_stack_lines:
            comprehensive_content += "\n".join(tech_stack_lines)
        else:
            # Fallback to a general message about retrieving from documents
            comprehensive_content += "Technology stack information will be retrieved from the provided documents based on the specific service requirements.\nPlease refer to the retrieved context for detailed technology stack information."
        
        return comprehensive_content
    else:
        # Generate a generic service description that works for any service
        service_description = f"{service_name} is a specialized service that involves creating custom solutions to meet specific business requirements. This service includes planning, development, testing, and deployment of tailored solutions to address unique challenges."
        
        # Extract technology stack information from the original response (RAG results)
        # Look for technology stack related content in the original response
        import re
        tech_stack_lines = []
        
        # Look for common tech stack patterns in the original response
        lines = original_response.split('\n')
        for line in lines:
            line = line.strip()
            # Look for lines that might contain tech stack information
            if any(keyword in line.lower() for keyword in ['frontend', 'backend', 'database', 'cloud', 'framework', 'library', 'technology', 'tool', 'platform', 'api', 'development', 'programming', 'programming language', 'javascript', 'python', 'java', 'react', 'angular', 'node', 'mysql', 'mongodb', 'aws', 'azure', 'docker', 'kubernetes']):
                if len(line) > 10 and not line.startswith('Regarding') and not line.startswith('Company') and not line.startswith('Cost'):
                    tech_stack_lines.append(line)
        
        # Generate the comprehensive format in conversational style with clear sections and proper spacing
        comprehensive_content = f"Company:\n\nSyngrid is a software development company based in Madurai, India, operating from the TBI@TCE at Thiruparankundaram. The company specializes in a range of digital services and can be contacted at (+91) 89250-31568.\n\n\nService Overview:\n\nRegarding {service_name} specifically:\n\n{service_description}\n\n\nCost:\n\nFor this service, Syngrid offers a fixed cost of {price}, which covers the complete development lifecycle from initial concept through final delivery.\n\n\nTechnology Stack:\n\n"
        
        if tech_stack_lines:
            comprehensive_content += "\n".join(tech_stack_lines)
        else:
            # Fallback to a general message about retrieving from documents
            comprehensive_content += "Technology stack information will be retrieved from the provided documents based on the specific service requirements.\nPlease refer to the retrieved context for detailed technology stack information."
        
        return comprehensive_content

def markdown_to_pdf(content: str, title: str = "Chat Response", is_quotation: bool = False) -> io.BytesIO:
    from fpdf import FPDF
    import os
    
    # Clean the content by removing UI elements like edit/download controls
    # Remove common UI elements that shouldn't appear in documents
    lines = content.split('\n')
    cleaned_lines = []
    
    for line in lines:
        if not any(ui_element in line for ui_element in ['✏️ Edit', '📄 PDF', '📝 Text', '📝 Word', 'Edit', 'Download', 'PDF', 'DOCX', 'TXT']):
            cleaned_lines.append(line)
    
    content = '\n'.join(cleaned_lines)
    
    # Replace rupee symbol with Rs. for PDF compatibility
    content = content.replace('₹', 'Rs.')
    
    class QuotationPDF(FPDF):
        def __init__(self):
            super().__init__()
            self.is_quotation = is_quotation
        
        def header(self):
            if self.is_quotation:
                # Use full-width header image
                header_path = os.path.join(os.path.dirname(__file__), "syngrid_header.png")
                if os.path.exists(header_path):
                    try:
                        # Add full-width header image (A4 width = 210mm)
                        self.image(header_path, x=0, y=0, w=210)
                        # Move down after header
                        self.ln(50)
                    except Exception as e:
                        logger.error(f"Failed to load header image: {e}")
                        # Fallback to simple header
                        self.set_fill_color(52, 116, 186)
                        self.rect(0, 0, 210, 45, 'F')
                        self.set_font('Helvetica', 'B', 20)
                        self.set_text_color(255, 255, 255)
                        self.set_xy(15, 15)
                        self.cell(0, 10, 'SYNGRID', 0, 0, 'L')
                        self.ln(20)
                else:
                    # Fallback if image not found
                    self.set_fill_color(52, 116, 186)
                    self.rect(0, 0, 210, 45, 'F')
                    self.set_font('Helvetica', 'B', 20)
                    self.set_text_color(255, 255, 255)
                    self.set_xy(15, 15)
                    self.cell(0, 10, 'SYNGRID', 0, 0, 'L')
                    self.ln(20)
            else:
                # Simple header for non-quotation
                self.set_font('Helvetica', 'B', 14)
                self.cell(0, 10, title, new_x='LMARGIN', new_y='NEXT', align='C')
                self.ln(5)
        
        def footer(self):
            if self.is_quotation:
                # Use full-width footer image
                footer_path = os.path.join(os.path.dirname(__file__), "syngrid_footer.png")
                if os.path.exists(footer_path):
                    try:
                        # Position footer at bottom (A4 height = 297mm, footer height ~15mm)
                        self.set_y(-15)
                        # Add full-width footer image
                        self.image(footer_path, x=0, y=self.get_y(), w=210)
                    except Exception as e:
                        logger.error(f"Failed to load footer image: {e}")
                        # Fallback to simple footer
                        self.set_y(-20)
                        self.set_fill_color(52, 73, 94)
                        self.rect(0, self.get_y(), 210, 20, 'F')
                        self.set_font('Helvetica', '', 11)
                        self.set_text_color(255, 255, 255)
                        self.set_y(-14)
                        self.cell(0, 10, 'syngrid.com', 0, 0, 'C')
                else:
                    # Fallback if image not found
                    self.set_y(-20)
                    self.set_fill_color(52, 73, 94)
                    self.rect(0, self.get_y(), 210, 20, 'F')
                    self.set_font('Helvetica', '', 11)
                    self.set_text_color(255, 255, 255)
                    self.set_y(-14)
                    self.cell(0, 10, 'syngrid.com', 0, 0, 'C')
            else:
                # Simple footer
                self.set_y(-15)
                self.set_font('Helvetica', 'I', 8)
                self.cell(0, 10, f'Page {self.page_no()}', new_x='RIGHT', new_y='TOP', align='C')
    
    pdf = QuotationPDF()
    pdf.is_quotation = is_quotation
    pdf.add_page()
    
    # Set margins - larger top/bottom for quotations to accommodate header/footer images
    if is_quotation:
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.set_top_margin(55)  # Space for header image
    else:
        pdf.set_auto_page_break(auto=True, margin=25)
    
    # Add quotation title if it's a quotation
    if is_quotation:
        pdf.set_font('Helvetica', 'B', 18)
        pdf.set_text_color(52, 116, 186)  # Syngrid blue
        pdf.cell(0, 12, 'PROJECT QUOTATION', 0, 1, 'C')
        pdf.ln(8)
    
    # Check if it's a table
    if '|' in content or '\t' in content:
        lines = content.strip().split('\n')
        table_rows = []
        
        for line in lines:
            if '|' in line and '---' not in line:
                cells = [c.strip() for c in line.strip('|').split('|')]
                table_rows.append(cells)
            elif '\t' in line:
                cells = [c.strip() for c in line.split('\t')]
                table_rows.append(cells)
        
        if table_rows:
            num_cols = len(table_rows[0])
            
            # Calculate column widths
            if num_cols == 4:
                # Category, Item, Price, Quantity
                col_widths = [45, 75, 35, 25]
            else:
                col_width = (pdf.w - 30) / num_cols
                col_widths = [col_width] * num_cols
            
            # Header row with Syngrid blue
            pdf.set_font('Helvetica', 'B', 11)
            pdf.set_fill_color(52, 116, 186)  # Syngrid blue
            pdf.set_text_color(255, 255, 255)
            for i, cell in enumerate(table_rows[0]):
                pdf.cell(col_widths[i], 10, cell[:40], border=1, new_x='RIGHT', new_y='TOP', align='C', fill=True)
            pdf.ln()
            
            # Data rows
            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(0, 0, 0)
            for row_idx, row in enumerate(table_rows[1:]):
                # Alternate row colors
                if row_idx % 2 == 0:
                    pdf.set_fill_color(245, 245, 245)
                else:
                    pdf.set_fill_color(255, 255, 255)
                
                # Check if this is a total row
                is_total = row[0].lower().strip() in ['total', 'grand total', 'subtotal']
                if is_total:
                    pdf.set_font('Helvetica', 'B', 11)
                    pdf.set_fill_color(52, 116, 186)  # Syngrid blue for total
                    pdf.set_text_color(255, 255, 255)
                
                for i, cell in enumerate(row):
                    align = 'R' if i >= 2 else 'L'  # Right align price and quantity
                    pdf.cell(col_widths[i], 9, cell[:40], border=1, new_x='RIGHT', new_y='TOP', align=align, fill=True)
                
                pdf.ln()
                
                if is_total:
                    pdf.set_font('Helvetica', '', 10)
                    pdf.set_text_color(0, 0, 0)
    else:
        # Handle conversational format content with proper formatting
        # Identify and format specific sections differently
        lines = content.split('\n')
        
        for line in lines:
            stripped_line = line.strip()
            
            if stripped_line:
                # Check if this line is a section header (contains colons or is in all caps)
                if ':' in stripped_line and not stripped_line.startswith('Rs.') and not stripped_line.startswith('₹'):
                    # This looks like a section header (e.g., "Company:", "Service Overview:", "Cost:")
                    pdf.set_font('Helvetica', 'B', 12)  # Bold for headers
                    pdf.set_text_color(52, 116, 186)  # Syngrid blue
                    page_width = pdf.w - 2*pdf.l_margin  # Subtract left and right margins
                    pdf.multi_cell(page_width, 7, stripped_line, align='L')
                    pdf.ln(1)  # Add small space after header
                    pdf.set_font('Helvetica', '', 10)
                    pdf.set_text_color(0, 0, 0)  # Reset to black
                elif 'regarding' in stripped_line.lower() and 'specifically:' in stripped_line.lower():
                    # This is a service-specific header like "Regarding X specifically:"
                    pdf.set_font('Helvetica', 'B', 11)  # Bold for service headers
                    pdf.set_text_color(0, 0, 0)  # Black
                    page_width = pdf.w - 2*pdf.l_margin  # Subtract left and right margins
                    pdf.multi_cell(page_width, 6, stripped_line, align='L')
                    pdf.ln(1)  # Add small space after header
                    pdf.set_font('Helvetica', '', 10)
                elif any(keyword in stripped_line.lower() for keyword in ['overview:', 'cost:', 'technology stack:', 'total cost:', 'service 1:', 'service 2:', 'service 3:', 'service 4:', 'service 5:']):
                    # Other service-specific headers
                    pdf.set_font('Helvetica', 'B', 11)  # Bold for service headers
                    pdf.set_text_color(0, 0, 0)  # Black
                    page_width = pdf.w - 2*pdf.l_margin  # Subtract left and right margins
                    pdf.multi_cell(page_width, 6, stripped_line, align='L')
                    pdf.ln(1)  # Add small space after header
                    pdf.set_font('Helvetica', '', 10)
                elif any(tech in stripped_line for tech in ['ReactJS', 'Angular', 'Vue.js', 'NodeJS', 'Python', 'Java', 'MySQL', 'MongoDB', 'AWS', 'Azure', 'Docker']):
                    # Format individual tech stack items with indentation
                    pdf.set_font('Helvetica', '', 10)
                    pdf.set_text_color(0, 0, 0)
                    # Add indentation with cell then multi_cell
                    pdf.cell(10, 5, '', new_x='RIGHT', new_y='TOP')  # Indent
                    page_width = pdf.w - 2*pdf.l_margin - 10  # Account for indent
                    pdf.multi_cell(page_width, 5, f'- {stripped_line}', align='L')
                    pdf.ln(0.5)  # Add small space after list item
                elif stripped_line.startswith('-') or stripped_line.startswith('•'):
                    # Format bullet points and list items
                    pdf.set_font('Helvetica', '', 10)
                    pdf.set_text_color(0, 0, 0)
                    # Add indentation with cell then multi_cell
                    pdf.cell(10, 5, '', new_x='RIGHT', new_y='TOP')  # Indent
                    page_width = pdf.w - 2*pdf.l_margin - 10  # Account for indent
                    pdf.multi_cell(page_width, 5, stripped_line.replace('•', '-'), align='L')
                    pdf.ln(0.5)  # Add small space after list item
                elif stripped_line.startswith('1.') or stripped_line.startswith('2.') or stripped_line.startswith('3.') or stripped_line.startswith('4.') or stripped_line.startswith('5.'):
                    # Format numbered lists
                    pdf.set_font('Helvetica', '', 10)
                    pdf.set_text_color(0, 0, 0)
                    # Add indentation with cell then multi_cell
                    pdf.cell(10, 5, '', new_x='RIGHT', new_y='TOP')  # Indent
                    page_width = pdf.w - 2*pdf.l_margin - 10  # Account for indent
                    pdf.multi_cell(page_width, 5, stripped_line, align='L')
                    pdf.ln(0.5)  # Add small space after list item
                elif len(stripped_line) > 2 and stripped_line.isupper():
                    # Format all-caps headers
                    pdf.set_font('Helvetica', 'B', 11)  # Bold for headers
                    pdf.set_text_color(0, 0, 0)  # Black
                    page_width = pdf.w - 2*pdf.l_margin  # Subtract left and right margins
                    pdf.multi_cell(page_width, 6, stripped_line, align='L')
                    pdf.ln(1)  # Add small space after header
                    pdf.set_font('Helvetica', '', 10)
                else:
                    # Regular text
                    pdf.set_font('Helvetica', '', 10)
                    pdf.set_text_color(0, 0, 0)
                    # Use the page width minus margins for the text
                    page_width = pdf.w - 2*pdf.l_margin  # Subtract left and right margins
                    pdf.multi_cell(page_width, 5, stripped_line, align='L')
                    pdf.ln(0.5)  # Add small space after regular text
            else:
                # Empty line - add minimal space for paragraph breaks
                pdf.ln(2)  # Reduced space for paragraph breaks
    
    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf

def markdown_to_docx(content: str, title: str = "Chat Response", is_quotation: bool = None) -> io.BytesIO:
    content = content.replace('₹', 'Rs.')
    
    # Clean the content by removing UI elements like edit/download controls
    # Remove common UI elements that shouldn't appear in documents
    lines = content.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # Skip lines that contain UI controls like edit/download icons
        if not any(ui_element in line for ui_element in ['✏️ Edit', '📄 PDF', '📝 Text', '📝 Word', 'Edit', 'Download', 'PDF', 'DOCX', 'TXT']):
            cleaned_lines.append(line)
    
    content = '\n'.join(cleaned_lines)
    
    doc = DocxDocument()
    
    # Determine if this is a quotation based on parameter or title
    if is_quotation is None:
        is_quotation = "quotation" in title.lower() or "Quotation" in title
    
    # Add header with logo if it's a quotation
    if is_quotation:
        # Try to add header with image
        try:
            header = doc.sections[0].header
            header_paragraph = header.paragraphs[0]
            header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Try to add header image
            header_image_path = os.path.join(os.path.dirname(__file__), "syngrid_header.png")
            if os.path.exists(header_image_path):
                # Add image to header (scaled appropriately for DOCX)
                header_paragraph.add_run().add_picture(header_image_path, width=Inches(5))
            else:
                # Fallback to text header
                header_run = header_paragraph.add_run("SYNGRID TECHNOLOGIES\n")
                header_run.bold = True
                header_run.font.size = Pt(16)
        except Exception as e:
            logger.error(f"Failed to add header to DOCX: {e}")
    
    # Add title
    title_para = doc.add_heading(title, 0)
    if is_quotation:
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].font.size = Pt(18)
        title_para.runs[0].font.color.rgb = RGBColor(52, 116, 186)  # Syngrid blue
    
    if ('|' in content and '---' in content) or ('\t' in content and 'Category' in content):
        lines = content.strip().split('\n')
        table_rows = []
        
        for line in lines:
            if '|' in line and '---' not in line:
                # Pipe separated
                cells = [c.strip() for c in line.strip('|').split('|')]
                table_rows.append(cells)
            elif '\t' in line:
                # Tab separated
                cells = [c.strip() for c in line.split('\t')]
                table_rows.append(cells)
        
        if table_rows:
            table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            table.style = 'Table Grid'
            
            # Style the table
            for i, row in enumerate(table_rows):
                for j, cell in enumerate(row):
                    table_cell = table.cell(i, j)
                    table_cell.text = cell
                    
                    # Style header row
                    if i == 0:
                        for paragraph in table_cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(11)
                    
                    # Style total row
                    if row[0].lower().startswith('total'):
                        for paragraph in table_cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(11)
                                run.font.color.rgb = RGBColor(52, 116, 186)  # Syngrid blue
            
            # Center align the table
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set column widths for better presentation
            widths = [Inches(2.5), Inches(3.5), Inches(1.5), Inches(1)]
            for i, width in enumerate(widths):
                for row in table.rows:
                    row.cells[i].width = width
    else:
        doc.add_paragraph(content)
    
    # Add footer with page numbers if it's a quotation
    if is_quotation:
        try:
            # Add page number to footer
            footer = doc.sections[0].footer
            
            # Try to add footer image
            footer_image_path = os.path.join(os.path.dirname(__file__), "syngrid_footer.png")
            if os.path.exists(footer_image_path):
                # Clear existing paragraphs and add image
                for paragraph in footer.paragraphs:
                    p = paragraph._element
                    p.getparent().remove(p)
                
                # Add a new paragraph for the image
                footer_para = footer.add_paragraph()
                footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                footer_para.add_run().add_picture(footer_image_path, width=Inches(5))
            else:
                # Fallback to simple footer text
                if footer.paragraphs:
                    footer_paragraph = footer.paragraphs[0]
                else:
                    footer_paragraph = footer.add_paragraph()
                footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                footer_run = footer_paragraph.add_run("Page ")
                footer_run.add_page_number()
                footer_run.font.size = Pt(10)
        except Exception as e:
            logger.error(f"Failed to add footer to DOCX: {e}")
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ============= LLM & EMBEDDINGS =============
class OpenRouterLLM(LLM):
    api_key: str
    model: str
    temperature: float = 0.7
    max_tokens: int = 2000
    
    @property
    def _llm_type(self) -> str:
        return "openrouter"
    
    def _call(self, prompt: str, stop: Optional[List[str]] = None, **kwargs: Any) -> str:
        if not self.api_key:
            raise ValueError("OPENROUTER_API_KEY not set")
        
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }

        # Include a brief system message to help steer responses.
        messages = [
            {"role": "system", "content": "You are a helpful assistant that follows instructions precisely."},
            {"role": "user", "content": prompt}
        ]

        payload = {
            "model": self.model,
            "messages": messages,
            "temperature": self.temperature,
            "max_tokens": self.max_tokens
        }

        try:
            response = requests.post(OPENROUTER_API_BASE, headers=headers, json=payload, timeout=60)
            # Log status and a short snippet of response for diagnostics
            logger.info(f"OpenRouter status: {response.status_code}")

            # Try to parse JSON safely and support multiple response shapes
            try:
                data = response.json()
            except Exception:
                logger.error("OpenRouter response is not valid JSON; returning raw text")
                response.raise_for_status()
                return response.text

            if isinstance(data, dict):
                # choices -> message -> content
                choices = data.get("choices") or data.get("outputs") or []
                if choices and isinstance(choices, list):
                    first = choices[0]
                    # message.content
                    if isinstance(first, dict):
                        msg = first.get("message") or first.get("output") or first
                        if isinstance(msg, dict):
                            # content may be nested
                            content = msg.get("content") or msg.get("text") or msg.get("response")
                            if isinstance(content, str):
                                return content
                        # fallback: first.get('text')
                        if isinstance(first.get("text"), str):
                            return first.get("text")

            # As a last resort, return the raw response text
            logger.warning(f"Unexpected OpenRouter response shape: {json.dumps(data)[:1000]}")
            response.raise_for_status()
            return response.text
        except Exception as e:
            logger.error(f"OpenRouter API error: {e}", exc_info=True)
            # Return a helpful fallback message instead of completely failing
            if "503" in str(e) or "Service Unavailable" in str(e):
                return "I'm sorry, but the AI service is currently unavailable. Please try again later. In the meantime, I can help with general information about our services."
            else:
                raise

class SentenceTransformerEmbeddings(Embeddings):
    def __init__(self, model_name: str = "all-MiniLM-L12-v2"):
        # Ensure compatibility with huggingface_hub versions that removed
        # `cached_download`. Some versions of `sentence-transformers` still
        # import `cached_download` from `huggingface_hub` at import time.
        try:
            import huggingface_hub
            if not hasattr(huggingface_hub, "cached_download"):
                from huggingface_hub import hf_hub_download

                def _cached_download(*args, **kwargs):
                    # Provide a compatibility wrapper for older callers that
                    # pass a `url` kwarg (used by some older sentence-transformers).
                    # If `url` is provided, download it directly and return
                    # the local filepath. Otherwise forward to `hf_hub_download`.
                    url = kwargs.pop('url', None)
                    if url:
                        import tempfile
                        try:
                            import requests
                        except Exception:
                            raise RuntimeError("requests is required to download legacy URLs")

                        try:
                            resp = requests.get(url, stream=True, timeout=30)
                            resp.raise_for_status()
                            tmpf = tempfile.NamedTemporaryFile(delete=False)
                            with tmpf as f:
                                for chunk in resp.iter_content(chunk_size=8192):
                                    if chunk:
                                        f.write(chunk)
                            return tmpf.name
                        except Exception as e:
                            logger.error(f"Failed to download from url {url}: {e}")
                            raise

                    try:
                        return hf_hub_download(*args, **kwargs)
                    except TypeError:
                        try:
                            return hf_hub_download(*args)
                        except Exception:
                            raise

                huggingface_hub.cached_download = _cached_download
                logger.info("Patched huggingface_hub.cached_download -> hf_hub_download")
        except Exception as e:
            logger.debug(f"Could not patch huggingface_hub: {e}")

        from sentence_transformers import SentenceTransformer
        self.model = SentenceTransformer(model_name)
        self.model_name = model_name
        self.method = "SentenceTransformer"
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return self.model.encode(texts, show_progress_bar=True).tolist()
    
    def embed_query(self, text: str) -> List[float]:
        return self.model.encode([text])[0].tolist()

# ============= VECTOR DB =============
def load_all_documents() -> List[Document]:
    docs = []
    for filename in os.listdir(DATA_DIR):
        filepath = os.path.join(DATA_DIR, filename)
        if not os.path.isfile(filepath):
            continue
        
        ext = os.path.splitext(filename.lower())[1]
        if ext not in SUPPORTED_EXTENSIONS:
            continue
        
        try:
            loader = SUPPORTED_EXTENSIONS[ext](filepath)
            loaded = loader.load()
            for d in loaded:
                d.metadata = d.metadata or {}
                d.metadata["source"] = filename
            docs.extend(loaded)
            logger.info(f"Loaded {len(loaded)} docs from {filename}")
        except Exception as e:
            logger.error(f"Failed to load {filename}: {e}")
    
    return docs

def create_vector_db():
    try:
        docs = load_all_documents()
        if not docs:
            logger.warning("No documents found in data directory")
            return None
        
        logger.info(f"Loaded {len(docs)} documents, splitting into chunks...")
        splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=300)
        chunks = splitter.split_documents(docs)
        logger.info(f"Created {len(chunks)} chunks from documents")
        
        global embeddings
        logger.info("Initializing sentence transformer embeddings...")
        embeddings = SentenceTransformerEmbeddings()
        logger.info("Embeddings initialized successfully")
        
        logger.info("Creating FAISS vector database...")
        db = FAISS.from_documents(chunks, embeddings)
        
        # Ensure directory exists
        os.makedirs(FAISS_DB_PATH, exist_ok=True)
        
        logger.info(f"Saving vector database to {FAISS_DB_PATH}...")
        db.save_local(FAISS_DB_PATH)
        logger.info(f"Vector DB created successfully with {len(chunks)} chunks")
        return db
    except Exception as e:
        logger.error(f"Failed to create vector DB: {e}", exc_info=True)
        return None

def load_vector_db():
    try:
        if not os.path.exists(FAISS_DB_PATH):
            logger.info(f"Vector DB path does not exist: {FAISS_DB_PATH}")
            return None
        
        # Check if index file exists
        index_file = os.path.join(FAISS_DB_PATH, "index.faiss")
        if not os.path.exists(index_file):
            logger.info(f"FAISS index file not found at {index_file}")
            return None
        
        global embeddings
        logger.info("Loading embeddings for vector DB...")
        embeddings = SentenceTransformerEmbeddings()
        
        logger.info(f"Loading FAISS vector database from {FAISS_DB_PATH}...")
        try:
            # Try with allow_dangerous_deserialization for newer versions
            db = FAISS.load_local(FAISS_DB_PATH, embeddings, allow_dangerous_deserialization=True)
            logger.info("Vector DB loaded successfully")
            return db
        except TypeError:
            # Fall back to older version without the parameter
            db = FAISS.load_local(FAISS_DB_PATH, embeddings)
            logger.info("Vector DB loaded successfully (legacy mode)")
            return db
    except Exception as e:
        logger.error(f"Failed to load vector DB: {e}", exc_info=True)
        return None

class SimpleRetrievalQA:
    def __init__(self, llm, vector_db, k=10):
        self.llm = llm
        self.vector_db = vector_db
        self.k = k
    
    def invoke(self, inputs):
        query = inputs.get("query", "")
        docs = self.vector_db.similarity_search(query, k=self.k)
        
        context = "\n\n".join([f"[{d.metadata.get('source', 'unknown')}]\n{d.page_content}" for d in docs])
        
        prompt = f"""You are a helpful assistant for Syngrid company. Answer based on this context.

RESPONSE INSTRUCTIONS:
- NEVER use table formats (no markdown tables with | separators)
- NEVER use markdown formatting (no **, *, _, etc.)
- For quotations/pricing requests, respond in conversational format with clear sections: Company, Service Overview, Cost, Technology Stack
- For explanation requests about specific services, provide detailed descriptions including features, benefits, and implementation details
- When explaining services, use natural language and be comprehensive
- Present information in clean, readable paragraphs without any markdown styling
- For price/quantity requests, calculate totals by multiplying Price × Quantity
- Always use ₹ for Indian Rupees
- If user specifies custom prices or quantities, use those values

Context:
{context}

Question: {query}

Based on the context, provide a helpful and accurate response:"""
        
        result = self.llm._call(prompt)
        return {"result": result, "source_documents": docs}

def setup_qa_chain(db, llm_instance):
    if db:
        return SimpleRetrievalQA(llm_instance, db)
    else:
        # Simple fallback
        class SimpleLLM:
            def __init__(self, llm):
                self.llm = llm
            def run(self, question):
                return self.llm._call(f"Answer: {question}")
        return SimpleLLM(llm_instance)

def initialize_chatbot():
    global llm, qa_chain, vector_db
    
    _load_chat_history()
    
    llm = OpenRouterLLM(api_key=OPENROUTER_API_KEY, model=MODEL_NAME)
    
    try:
        vector_db = load_vector_db()
        if not vector_db:
            vector_db = create_vector_db()
    except Exception as e:
        logger.warning(f"Failed to initialize vector DB: {e}")
        logger.warning("App will run without vector database. Upload documents and rebuild to enable RAG.")
        vector_db = None
    
    qa_chain = setup_qa_chain(vector_db, llm)
    
    logger.info(f"Chatbot initialized - {len(chat_history)} chats loaded")

try:
    initialize_chatbot()
except Exception as e:
    logger.error(f"Initialization error: {e}")
    logger.info("Starting with minimal configuration...")

@app.route("/")
def home():
    return render_template("index.html") if os.path.exists("templates/index.html") else "<h3>RAG Chatbot Running</h3>"

@app.route("/chat")
def chat():
    return render_template("chat.html") if os.path.exists("templates/chat.html") else "<h3>Chat interface not found</h3>"

@app.route("/api/health")
def health():
    return jsonify({"status": "healthy", "timestamp": datetime.utcnow().isoformat()})

@app.route("/api/status")
def status():
    files = [f for f in os.listdir(DATA_DIR) if os.path.splitext(f.lower())[1] in SUPPORTED_EXTENSIONS]
    return jsonify({
        "vector_db": vector_db is not None,
        "file_count": len(files),
        "chat_history_count": len(chat_history),
        "embedding_model": embeddings.model_name if embeddings else None
    })

@app.route("/api/chat", methods=["POST"])
def chat_api():
    try:
        data = request.get_json() or {}
        msg = data.get("message", "").strip()
        chat_id = data.get("chat_id")  # Get the chat_id if provided
        
        if not msg:
            return jsonify({"error": "Empty message"}), 400
        
        user_lower = msg.lower()
        
        # Detect quotation requests
        is_quotation_request = "quotation" in user_lower and ("create" in user_lower or "generate" in user_lower or "make" in user_lower) and ("for" in user_lower)
        
        # Detect price and quantity requests
        has_price_and_quantity = ("price" in user_lower and ("quantity" in user_lower or "qty" in user_lower or any(str(q) in user_lower for q in range(1, 20))))
        has_service_indicators = any(service in user_lower for service in ["chatbot", "website", "mobile app", "ai", "development", "design"])
        # Also consider it a price/quantity request if it contains price and quantity even without explicit service indicators
        is_price_quantity_request = has_price_and_quantity
        
        # No table generation anymore
        wants_table = False
        
        # Determine if this should be processed as a quotation
        should_process_as_quotation = is_quotation_request or is_price_quantity_request
        
        # Enhanced prompt (no table instructions unless explicitly requested)
        query = msg
        
        conversation_context = ""
        if chat_id and chat_id in chat_history:
            chat_entry = chat_history[chat_id]
            if "messages" in chat_entry:
                # Limit the conversation context to the last 10 messages to avoid token limits
                recent_messages = chat_entry["messages"][-10:]
                for message in recent_messages:
                    role = message["role"].capitalize()
                    content = message["content"]
                    conversation_context += f"{role}: {content}\n"
            elif "bot_responses" in chat_entry:
                # Handle legacy format
                # Limit the conversation context to the last 5 bot responses
                recent_responses = chat_entry["bot_responses"][-5:]
                for i, response in enumerate(recent_responses):
                    conversation_context += f"Assistant: {response}\n"
            
            # Add the current query to the context
            full_query = conversation_context + f"User: {query}\nAssistant: "
        else:
            # For new chats, just use the current query
            full_query = f"User: {query}\nAssistant: "
        
        # Get response from LLM
        if vector_db:
            result = qa_chain.invoke({"query": full_query})
            answer = result["result"]
        else:
            answer = qa_chain.run(question=full_query)
        
        if should_process_as_quotation:
            # Handle quotation and price/quantity requests in conversational format
            
            # Extract multiple service names from the user message
            # Look for patterns like "for iOS Development", "for website", etc.
            import re
            
            # Find all services mentioned in the message
            services_found = []
            
            # Pattern 1: "for <service>" - find the service(s)
            for_match = re.search(r'for\s+(.+?)(?:$)', msg, re.IGNORECASE)
            if for_match:
                services_text = for_match.group(1).strip()
                
                # Split by 'and', '&', or ',' to separate multiple services, but avoid splitting compound terms
                # Check if the service text contains clear multiple service indicators
                # but avoid splitting within compound service names like "UI/UX", "AI Chatbot", etc.
                
                # First, identify if we have multiple services by looking for 'and', '&', or ','
                # but make sure these are not part of compound terms
                service_parts = [services_text]  # Start with the full text as one service
                
                # Check if we have clear multiple service indicators that are not part of compound terms
                # Split only if 'and', '&', or ',' are not within known compound terms
                if re.search(r'\s+and\s+|\s*&\s+|,\s*', services_text, re.IGNORECASE):
                    # Check for potential compound terms that should not be split
                    potential_split = re.split(r'\s+and\s+|\s*&\s*|,\s*', services_text)
                    
                    # Check if any part contains known compound terms that shouldn't be split
                    compound_indicators = ['ui/ux', 'ai chatbot', 'payment gateway', 'api integration']
                    has_compound = any(compound in services_text.lower() for compound in compound_indicators)
                    
                    # Also check if we have potential multiple distinct services
                    distinct_services = []
                    potential_multiple = False
                    
                    for part in potential_split:
                        part_clean = part.strip().lower()
                        # Check if this part could be a distinct service by looking for service keywords
                        service_keywords = ['development', 'design', 'app', 'website', 'chatbot', 'api', 'gateway', 'integration']
                        if any(keyword in part_clean for keyword in service_keywords):
                            distinct_services.append(part.strip())
                            potential_multiple = True
                    
                    # Only split if we have multiple distinct services, not just compound terms
                    if potential_multiple and len(distinct_services) > 1:
                        service_parts = distinct_services
                    else:
                        service_parts = [services_text]  # Treat as single service
                
                for service_part in service_parts:
                    service = service_part.strip()
                    # Clean up any trailing words after service name
                    service = re.sub(r'\s+(?:with|with\s+|quotation|request)', '', service, flags=re.IGNORECASE)
                    
                    if service:  # Only add non-empty services
                        # Enhance service name if it's a common short form
                        service_lower = service.lower().strip()
                        if service_lower == "ios":
                            service = "iOS Development"
                        elif service_lower == "android":
                            service = "Android Development"
                        elif service_lower == "web":
                            service = "Website Development"
                        elif 'ui/ux' in service_lower or ('ui' in service_lower and 'ux' in service_lower) or ('design' in service_lower and ('ui' in service_lower or 'ux' in service_lower)):
                            # Handle UI/UX design services specifically
                            if not service_lower.endswith(('design', 'development', 'service', 'app', 'bot', 'system', 'solution', 'project')):
                                service = service.strip().title() + " Design"
                        elif service_lower == "mobile":
                            service = "Mobile App Development"
                        else:
                            # General handling for other services
                            if not service_lower.endswith(('development', 'service', 'app', 'bot', 'system', 'solution', 'project', 'design')):
                                # For design-related services
                                if 'design' in service_lower or 'ui' in service_lower or 'ux' in service_lower:
                                    service = service.strip().title() + " Design"
                                # For development-related services
                                elif 'development' in service_lower or 'dev' in service_lower:
                                    service = service.strip().title() + " Development"
                                # For service-related terms
                                elif 'service' in service_lower or 'consulting' in service_lower:
                                    service = service.strip().title() + " Service"
                                # Default to Development for most other cases
                                else:
                                    service = service.strip().title() + " Development"
                        services_found.append(service)
            

            

            
            # Fallback: try to find service keywords in the message
            service_keywords = ["ai chatbot", "website", "mobile app", "ios development", "android development", "ui/ux design", "api integration", "testing", "maintenance", "payment gateway"]
            for service in service_keywords:
                if service in user_lower:
                    # Check if this service is not already in the list
                    if not any(service in existing_service.lower() for existing_service in services_found):
                        services_found.append(service)
            
            # Remove duplicates while preserving order
            seen = set()
            unique_services = []
            for service in services_found:
                service_lower = service.lower()
                if service_lower not in seen:
                    seen.add(service_lower)
                    unique_services.append(service)
            
            # If no services were found, try to extract from original logic as fallback
            if not unique_services:
                # Fallback to original single-service extraction
                service_name = msg  # Use the entire message as service name if no specific pattern matches
                
                # Pattern 1: "for <service>"
                for_pattern = re.search(r'for\s+([a-zA-Z\s]+?)(?:\s+|$)', msg, re.IGNORECASE)
                if for_pattern:
                    service_name = for_pattern.group(1).strip()
                    
                    # Enhance service name if it's a common short form
                    service_name_lower = service_name.lower().strip()
                    if service_name_lower == "ios":
                        service_name = "iOS Development"
                    elif service_name_lower == "android":
                        service_name = "Android Development"
                    elif service_name_lower == "web":
                        service_name = "Website Development"
                    elif 'ui/ux' in service_name_lower or 'ui' in service_name_lower and 'ux' in service_name_lower or 'design' in service_name_lower and ('ui' in service_name_lower or 'ux' in service_name_lower):
                        # Handle UI/UX design services specifically
                        if not service_name_lower.endswith(('design', 'development', 'service', 'app', 'bot', 'system', 'solution', 'project')):
                            service_name = service_name.strip().title() + " Design"
                    elif service_name_lower == "mobile":
                        service_name = "Mobile App Development"
                else:
                            # Check if this is a follow-up request that references a previous service
                            # Look for previous service in conversation history
                            if chat_id and chat_id in chat_history:
                                chat_entry = chat_history[chat_id]
                                if "messages" in chat_entry:
                                    # Look through previous assistant messages to find the last mentioned service
                                    for message in reversed(chat_entry["messages"]):
                                        if message["role"] == "assistant":
                                            assistant_content = message["content"]
                                            # Look for service-specific patterns in the assistant's previous response
                                            if "Regarding" in assistant_content and "specifically:" in assistant_content:
                                                # Extract service from previous response format
                                                prev_service_match = re.search(r'Regarding\s+([a-zA-Z\s]+?)\s+specifically:', assistant_content)
                                                if prev_service_match:
                                                    prev_service = prev_service_match.group(1).strip()
                                                    # Only use if it looks like a real service name (not just 'service')
                                                    if prev_service.lower() != "service" and len(prev_service) > 1:
                                                        service_name = prev_service
                                                        break
                            
                            # Fallback: try to find service keywords in the message
                            service_keywords = ["ai chatbot", "website", "mobile app", "ios development", "android development", "ui/ux design", "api integration", "testing", "maintenance"]
                            for service in service_keywords:
                                if service in user_lower:
                                    service_name = service
                                    break
                            
                            # If still no specific service found, use a more descriptive default based on the message
                            if service_name.lower().strip() == msg.lower().strip():
                                # Extract the main subject from the message if it's still the full message
                                service_name = msg.replace("create quotation for", "").replace("generate quotation for", "").replace("make quotation for", "").strip()
                                if not service_name or service_name.lower() == "service":
                                    service_name = "Custom Development Service"  # More descriptive fallback
                
                # Generate comprehensive format without table for single service
                answer = generate_comprehensive_format(service_name, answer, msg)
            else:
                # Multiple services found, generate comprehensive format for all services
                answer = generate_comprehensive_format(unique_services, answer, msg)
        else:
            # For all other responses, remove markdown formatting
            answer = remove_markdown_formatting(answer)
            

        
        # Always set is_table to False since we're not generating tables anymore
        wants_table = False
        
        # Save chat - if chat_id provided, update existing chat; otherwise create new one
        if chat_id and chat_id in chat_history:
            # Update existing chat - add new bot response to the list
            chat_entry = update_chat(chat_id, bot_response=answer)
        else:
            # Create new chat
            chat_entry = create_chat(msg, answer, wants_table)
        
        return jsonify({
            "response": answer,
            "chat_id": chat_entry["id"] if chat_entry else None,
            "is_table": wants_table
        })
        
    except Exception as e:
        logger.error(f"Chat error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500

@app.route("/api/chats", methods=["GET"])
def list_chats():
    limit = request.args.get("limit", 100, type=int)
    chats = list(chat_history.values())
    chats.sort(key=lambda x: x["created_at"], reverse=True)
    return jsonify(chats[:limit])

@app.route("/api/chats/<chat_id>", methods=["GET"])
def get_chat(chat_id):
    chat = chat_history.get(chat_id)
    if not chat:
        return jsonify({"error": "Not found"}), 404
    
    # For backward compatibility, return the most recent bot response as bot_response
    response_chat = chat.copy()
    
    # Get the most recent bot response from messages
    if "messages" in chat and chat["messages"]:
        # Find the last assistant message
        bot_messages = [msg for msg in chat["messages"] if msg["role"] == "assistant"]
        if bot_messages:
            response_chat["bot_response"] = bot_messages[-1]["content"]
        else:
            response_chat["bot_response"] = ""
        
        # Get the last user message
        user_messages = [msg for msg in chat["messages"] if msg["role"] == "user"]
        if user_messages:
            response_chat["user_message"] = user_messages[-1]["content"]
        else:
            response_chat["user_message"] = ""
    elif "bot_responses" in chat and chat["bot_responses"]:
        response_chat["bot_response"] = chat["bot_responses"][-1]
        response_chat["user_message"] = chat.get("user_message", "")
    elif "bot_response" not in chat and "bot_responses" not in chat:
        response_chat["bot_response"] = ""
        response_chat["user_message"] = ""
    
    return jsonify(response_chat)

@app.route("/api/chats/<chat_id>", methods=["PUT"])
def update_chat_route(chat_id):
    data = request.get_json() or {}
    updated = update_chat(chat_id, bot_response=data.get("bot_response"), user_message=data.get("user_message"))
    return jsonify(updated) if updated else (jsonify({"error": "Not found"}), 404)

@app.route("/api/chats/<chat_id>", methods=["DELETE"])
def delete_chat_route(chat_id):
    success = delete_chat(chat_id)
    return jsonify({"message": "Deleted"}) if success else (jsonify({"error": "Not found"}), 404)

@app.route("/api/chats/<chat_id>/download", methods=["GET"])
def download_chat(chat_id):
    chat = chat_history.get(chat_id)
    if not chat:
        return jsonify({"error": "Not found"}), 404
    
    fmt = request.args.get("format", "txt")
    
    # Detect if this is a quotation
    user_msg = chat.get("user_message", "").lower()
    is_quotation = any(keyword in user_msg for keyword in ["quotation", "quote", "pricing", "price", "cost"])
    
    # Get the most recent bot response
    bot_response = chat.get("bot_response", "")  # For backward compatibility
    
    # Handle new message structure
    if "messages" in chat and chat["messages"]:
        # Find the last assistant message
        bot_messages = [msg for msg in chat["messages"] if msg["role"] == "assistant"]
        if bot_messages:
            bot_response = bot_messages[-1]["content"]
        
        # Update user_msg from the messages if needed
        user_messages = [msg for msg in chat["messages"] if msg["role"] == "user"]
        if user_messages:
            user_msg = user_messages[-1]["content"].lower()
            is_quotation = any(keyword in user_msg for keyword in ["quotation", "quote", "pricing", "price", "cost"])
    elif "bot_responses" in chat and chat["bot_responses"]:
        bot_response = chat["bot_responses"][-1]  # Get the most recent response
    elif "bot_response" in chat:  
        bot_response = chat["bot_response"]
    
    if fmt == "csv" and chat.get("is_table"):
        buf = io.BytesIO(markdown_table_to_csv(bot_response).encode())
        return send_file(buf, as_attachment=True, download_name=f"quotation_{chat_id[:8]}.csv" if is_quotation else f"chat_{chat_id[:8]}.csv", mimetype="text/csv")
    
    elif fmt == "xlsx" and chat.get("is_table"):
        buf = markdown_table_to_excel(bot_response)
        return send_file(buf, as_attachment=True, download_name=f"quotation_{chat_id[:8]}.xlsx" if is_quotation else f"chat_{chat_id[:8]}.xlsx", 
                        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    
    elif fmt == "pdf":
        # Use professional quotation format if it's a quotation request
        buf = markdown_to_pdf(bot_response, "Quotation" if is_quotation else "Chat Response", is_quotation=is_quotation)
        download_name = f"Syngrid_Quotation_{chat_id[:8]}.pdf" if is_quotation else f"chat_{chat_id[:8]}.pdf"
        return send_file(buf, as_attachment=True, download_name=download_name, mimetype="application/pdf")
    
    elif fmt == "docx":
        buf = markdown_to_docx(bot_response, "Quotation" if is_quotation else "Chat Response", is_quotation=is_quotation)
        return send_file(buf, as_attachment=True, download_name=f"quotation_{chat_id[:8]}.docx" if is_quotation else f"chat_{chat_id[:8]}.docx",
                        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    else:  # txt
        # For table data (especially quotations), show only the table content
        if chat.get("is_table"):
            # Clean the content by removing extra prefixes
            content = bot_response.strip()
            # Remove any 'Bot:' prefix if present
            if content.startswith('Bot:'):
                content = content[4:].strip()
        else:
            # For non-table content, keep the original format but remove markdown
            content = f"User: {chat.get('user_message', '')}\n\nBot: {remove_markdown_formatting(bot_response)}\n"
        buf = io.BytesIO(content.encode())
        # Use appropriate filename based on content type
        if chat.get("is_table"):
            download_name = f"quotation_{chat_id[:8]}.txt" if is_quotation else f"table_{chat_id[:8]}.txt"
        else:
            download_name = f"chat_{chat_id[:8]}.txt"
        return send_file(buf, as_attachment=True, download_name=download_name, mimetype="text/plain")

@app.route("/api/chats/clear", methods=["POST"])
def clear_chats():
    count = len(chat_history)
    for cid in list(chat_history.keys()):
        delete_chat(cid)
    return jsonify({"message": f"Deleted {count} chats", "deleted_count": count})

@app.route("/api/documents", methods=["GET"])
def list_documents():
    files = []
    for filename in os.listdir(DATA_DIR):
        filepath = os.path.join(DATA_DIR, filename)
        if os.path.isfile(filepath):
            stat = os.stat(filepath)
            ext = os.path.splitext(filename.lower())[1]
            files.append({
                "name": filename,
                "extension": ext,
                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                "supported": ext in SUPPORTED_EXTENSIONS,
                "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat()
            })
    return jsonify({"documents": files, "count": len(files)})

@app.route("/api/documents/upload", methods=["POST"])
def upload_document():
    try:
        logger.info(f"Upload request received. Files: {request.files}")
        logger.info(f"Form data: {request.form}")
        
        if "file" not in request.files:
            logger.error("No file in request.files")
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files["file"]
        logger.info(f"File object: {file}, filename: {file.filename}")
        
        if not file.filename:
            logger.error("Empty filename")
            return jsonify({"error": "No file selected"}), 400
        
        filename = secure_filename(file.filename)
        ext = os.path.splitext(filename.lower())[1]
        logger.info(f"Secure filename: {filename}, extension: {ext}")
        
        if ext not in SUPPORTED_EXTENSIONS:
            logger.error(f"Unsupported extension: {ext}")
            return jsonify({"error": f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS.keys())}"}), 400
        
        filepath = os.path.join(DATA_DIR, filename)
        if os.path.exists(filepath):
            name, ext = os.path.splitext(filename)
            filename = f"{name}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}{ext}"
            filepath = os.path.join(DATA_DIR, filename)
        
        logger.info(f"Saving file to: {filepath}")
        file.save(filepath)
        logger.info(f"File saved successfully: {filename}")
        
        return jsonify({"message": "Uploaded successfully", "filename": filename})
    except Exception as e:
        logger.error(f"Upload error: {e}", exc_info=True)
        return jsonify({"error": f"Upload failed: {str(e)}"}), 500

@app.route("/api/documents/<filename>", methods=["DELETE"])
def delete_document(filename):
    filepath = os.path.join(DATA_DIR, secure_filename(filename))
    if not os.path.exists(filepath):
        return jsonify({"error": "Not found"}), 404
    os.remove(filepath)
    return jsonify({"message": "Deleted"})

@app.route("/api/rebuild", methods=["POST"])
def rebuild():
    try:
        if os.path.exists(FAISS_DB_PATH):
            shutil.rmtree(FAISS_DB_PATH)
        
        global vector_db, qa_chain
        vector_db = create_vector_db()
        qa_chain = setup_qa_chain(vector_db, llm)
        
        return jsonify({"message": "Rebuilt successfully"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)